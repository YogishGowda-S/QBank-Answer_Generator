from flask import Flask, render_template, request, redirect, session, send_file
import sqlite3
import os
import re
import docx
import pdfplumber
from groq import Groq
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

app = Flask(__name__)
app.secret_key = "qbank_secret_key"

try:
    groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
except TypeError:
    # Workaround for Groq compatibility
    import httpx
    groq_client = Groq(
        api_key=os.getenv("GROQ_API_KEY"),
        http_client=httpx.Client()
    )

# SQLite Database
DATABASE = 'qbank_app.db'

def get_db():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS subjects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id INTEGER NOT NULL,
            question_text TEXT NOT NULL,
            marks INTEGER,
            unit TEXT,
            FOREIGN KEY (subject_id) REFERENCES subjects(id)
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS note_chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id INTEGER NOT NULL,
            source_file TEXT,
            chunk_text TEXT NOT NULL,
            FOREIGN KEY (subject_id) REFERENCES subjects(id)
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER NOT NULL,
            answer_text TEXT,
            word_count INTEGER,
            FOREIGN KEY (question_id) REFERENCES questions(id)
        )
    ''')
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

def parse_questions(docx_paths):
    questions = []
    seen = set()

    skip_keywords = [
        'subject name', 'subject code', 'question bank',
        'department', 'semester', 'module', 'unit', 'section',
        'part a', 'part b', 'sl no', 'sl.no', 'course code',
        'course name', 'faculty', 'year', 'batch', 'college',
        'ia1', 'ia2', 'ia3', 'internal assessment'
    ]

    for docx_path in docx_paths:
        doc = docx.Document(docx_path)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 15:
                continue
            if text in seen:
                continue
            if any(kw in text.lower() for kw in skip_keywords):
                continue
            if len(text.split()) < 5:
                continue
            seen.add(text)
            match = re.search(r'\b(\d+)\s*[Mm]arks?\b|\[(\d+)\]|\((\d+)\)|\b(\d+)[Mm]\b', text)
            marks = 0
            if match:
                marks = int(next(m for m in match.groups() if m is not None))
            questions.append({"text": text, "marks": marks})

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text or len(text) < 15:
                        continue
                    if text in seen:
                        continue
                    if any(kw in text.lower() for kw in skip_keywords):
                        continue
                    if len(text.split()) < 5:
                        continue
                    seen.add(text)
                    match = re.search(r'\b(\d+)\s*[Mm]arks?\b|\[(\d+)\]|\((\d+)\)|\b(\d+)[Mm]\b', text)
                    marks = 0
                    if match:
                        marks = int(next(m for m in match.groups() if m is not None))
                    questions.append({"text": text, "marks": marks})

    return questions

def parse_notes(pdf_paths):
    all_text = ""
    for pdf_path in pdf_paths:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text += text + "\n\n"
        except Exception as e:
            print(f"Error reading {pdf_path}: {e}")
    return all_text

def clean_answer(raw):
    raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL)
    raw = re.sub(r'<thinking>.*?</thinking>', '', raw, flags=re.DOTALL)

    real_start_patterns = [
        r'(?=\*\*Introduction\*\*)',
        r'(?=\*\*Definition\*\*)',
        r'(?=\*\*Normalization\*\*)',
        r'(?=Introduction\n)',
        r'(?=Definition\n)',
    ]
    for pattern in real_start_patterns:
        match = re.search(pattern, raw, re.IGNORECASE)
        if match:
            raw = raw[match.start():]
            break

    lines = raw.split('\n')
    clean_lines = []
    skip_line_patterns = [
        "here's a thinking", 'thinking process',
        'analyze user input', '- role:', '- task:',
        '- marks:', '- length:', '- structure:', '- source:',
        'constraint check', 'deconstruct requirements',
        'mental to text', 'line count', 'write content',
        'identify key requirements', 'draft the answer',
        'never skip', 'output only', 'start every',
        'required length', 'specific format',
        "student's notes", 'for 10 mark questions',
        'write in clear paragraphs', 'when comparing',
        '(html table)', 'html table',
    ]
    for line in lines:
        line_lower = line.lower().strip()
        if any(p in line_lower for p in skip_line_patterns):
            continue
        if re.match(r'^\d+\.\s+\*\*(Analyze|Identify|Deconstruct|Draft|Write|Line|Constraint|Mental)', line):
            continue
        clean_lines.append(line)

    raw = '\n'.join(clean_lines)

    raw = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', raw)
    raw = re.sub(r'\*(.*?)\*', r'<em>\1</em>', raw)

    lines = raw.split('\n')
    result_lines = []
    in_table = False
    table_rows = []

    for line in lines:
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
        else:
            if in_table:
                html_table = "<table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse; width:100%; margin:10px 0;'>"
                for idx, row in enumerate(table_rows):
                    if all(re.match(r'^[-:]+$', cell) for cell in row if cell):
                        continue
                    html_table += "<tr>"
                    tag = "th" if idx == 0 else "td"
                    for cell in row:
                        html_table += f"<{tag} style='padding:6px; border:1px solid #ccc;'>{cell}</{tag}>"
                    html_table += "</tr>"
                html_table += "</table>"
                result_lines.append(html_table)
                in_table = False
                table_rows = []
            result_lines.append(line)

    if in_table and table_rows:
        html_table = "<table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse; width:100%; margin:10px 0;'>"
        for idx, row in enumerate(table_rows):
            if all(re.match(r'^[-:]+$', cell) for cell in row if cell):
                continue
            html_table += "<tr>"
            tag = "th" if idx == 0 else "td"
            for cell in row:
                html_table += f"<{tag} style='padding:6px; border:1px solid #ccc;'>{cell}</{tag}>"
            html_table += "</tr>"
        html_table += "</table>"
        result_lines.append(html_table)

    raw = '\n'.join(result_lines)
    raw = raw.replace('\n\n', '<br><br>').replace('\n', '<br>')
    return raw.strip()

def generate_one_answer(question, notes_text):
    notes_snippet = notes_text[:2000] if len(notes_text) > 2000 else notes_text
    marks = question["marks"] if question["marks"] > 0 else 10

    if marks <= 2:
        length_guide = "Write 3-5 lines. Give a clear definition and one key point."
    elif marks <= 5:
        length_guide = "Write 8-12 lines. Include explanation and a relevant example."
    else:
        length_guide = "Write 20-30 lines. Include introduction, detailed explanation with points, examples, and conclusion."

    prompt = f"""Answer this exam question for an engineering student.

Reference notes:
{notes_snippet}

Question: {question['text']}
Marks: {marks}
{length_guide}

Rules:
- Write ONLY the answer. Start directly with content.
- Use markdown bold (**text**) for headings.
- If comparing two concepts, use a markdown table like:
  | Feature | A | B |
  |---------|---|---|
  | Point   | x | y |
- Do NOT write any thinking process or analysis steps.
- Just write the clean exam answer."""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": "You are an exam answer writer for engineering students. Write only the answer directly. No thinking process. No analysis. No preamble. Start immediately with the answer."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=1500,
            temperature=0.3
        )
        raw = response.choices[0].message.content.strip()
        return clean_answer(raw)
    except Exception as e:
        return f"Could not generate answer: {str(e)}"

def generate_all_answers(questions, notes_text):
    all_answers = []
    for i, question in enumerate(questions):
        print(f"Generating Q{i+1}/{len(questions)}: {question['text'][:50]}...")
        answer = generate_one_answer(question, notes_text)
        all_answers.append(answer)
        print(f"Q{i+1} done.")
    return all_answers

def generate_pdf(questions, answers, output_path, subject_name):
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    safe_subject = re.sub(r'[<>&"\']', '', subject_name)
    story.append(Paragraph(f"{safe_subject} - Question Bank Answers", styles['Title']))
    story.append(Spacer(1, 20))

    for i, (question, answer) in enumerate(zip(questions, answers)):
        marks = question['marks'] if question['marks'] > 0 else 10

        q_text = f"Q{i+1} ({marks} marks): {question['text']}"
        q_text = re.sub(r'[<>&]', '', q_text)

        try:
            story.append(Paragraph(q_text, styles['Heading3']))
        except Exception:
            story.append(Paragraph(f"Q{i+1} ({marks} marks)", styles['Heading3']))

        story.append(Spacer(1, 6))

        clean = re.sub(r'<table.*?</table>', '', answer, flags=re.DOTALL)
        clean = re.sub(r'<.*?>', ' ', clean)
        clean = clean.replace('&nbsp;', ' ')
        clean = re.sub(r'\s+', ' ', clean).strip()
        clean = clean.replace('&', '&amp;')
        clean = clean.replace('<', '&lt;')
        clean = clean.replace('>', '&gt;')

        if clean:
            try:
                story.append(Paragraph(clean, styles['Normal']))
            except Exception:
                story.append(Paragraph("Answer available in web view.", styles['Normal']))

        story.append(Spacer(1, 20))

    doc.build(story)

@app.route("/")
def home():
    if "user_id" not in session:
        return redirect("/login")
    return redirect("/dashboard")

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        db = get_db()
        cursor = db.cursor()
        try:
            cursor.execute(
                "INSERT INTO users (email, password_hash) VALUES (?, ?)",
                (email, password)
            )
            db.commit()
            return redirect("/login")
        except:
            return "Email already exists. <a href='/signup'>Try again</a>"
        finally:
            db.close()
    return render_template("signup.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        db = get_db()
        cursor = db.cursor()
        cursor.execute(
            "SELECT id FROM users WHERE email=? AND password_hash=?",
            (email, password)
        )
        user = cursor.fetchone()
        db.close()
        if user:
            session["user_id"] = user[0]
            return redirect("/dashboard")
        else:
            return "Invalid credentials. <a href='/login'>Try again</a>"
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")

@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect("/login")
    db = get_db()
    cursor = db.cursor()
    cursor.execute(
        "SELECT id, name, created_at FROM subjects WHERE user_id=?",
        (session["user_id"],)
    )
    subjects = cursor.fetchall()
    db.close()
    return render_template("dashboard.html", subjects=subjects)

@app.route("/add_subject", methods=["GET", "POST"])
def add_subject():
    if "user_id" not in session:
        return redirect("/login")
    if request.method == "POST":
        name = request.form["name"]
        db = get_db()
        cursor = db.cursor()
        cursor.execute(
            "INSERT INTO subjects (user_id, name) VALUES (?, ?)",
            (session["user_id"], name)
        )
        db.commit()
        db.close()
        return redirect("/dashboard")
    return render_template("add_subject.html")

@app.route("/subject/<int:subject_id>")
def subject_detail(subject_id):
    if "user_id" not in session:
        return redirect("/login")
    db = get_db()
    cursor = db.cursor()
    cursor.execute(
        "SELECT id, name FROM subjects WHERE id=? AND user_id=?",
        (subject_id, session["user_id"])
    )
    subject = cursor.fetchone()
    cursor.execute(
        """SELECT q.question_text, q.marks, a.answer_text
           FROM questions q
           LEFT JOIN answers a ON q.id = a.question_id
           WHERE q.subject_id=?""",
        (subject_id,)
    )
    qa_pairs = cursor.fetchall()
    db.close()
    if not subject:
        return "Subject not found"
    return render_template("subject.html", subject=subject, qa_pairs=qa_pairs)

@app.route("/upload/<int:subject_id>", methods=["GET", "POST"])
def upload(subject_id):
    if "user_id" not in session:
        return redirect("/login")

    if request.method == "POST":
        qbank_files = request.files.getlist("qbank")
        notes_files = request.files.getlist("notes")

        upload_folder = f"uploads/{subject_id}"
        os.makedirs(upload_folder, exist_ok=True)

        qbank_paths = []
        for i, f in enumerate(qbank_files):
            if f and f.filename:
                path = os.path.join(upload_folder, f"qbank_{i}.docx")
                f.save(path)
                qbank_paths.append(path)

        notes_paths = []
        for i, f in enumerate(notes_files):
            if f and f.filename:
                path = os.path.join(upload_folder, f"notes_{i}.pdf")
                f.save(path)
                notes_paths.append(path)

        questions = parse_questions(qbank_paths)
        notes_text = parse_notes(notes_paths)

        print(f"Found {len(questions)} questions")

        db = get_db()
        cursor = db.cursor()
        cursor.execute(
            "DELETE FROM answers WHERE question_id IN (SELECT id FROM questions WHERE subject_id=?)",
            (subject_id,)
        )
        cursor.execute("DELETE FROM questions WHERE subject_id=?", (subject_id,))
        cursor.execute("DELETE FROM note_chunks WHERE subject_id=?", (subject_id,))
        db.commit()

        cursor.execute(
            "INSERT INTO note_chunks (subject_id, source_file, chunk_text) VALUES (?, ?, ?)",
            (subject_id, "combined_notes", notes_text[:100000])
        )

        answers = generate_all_answers(questions, notes_text)

        for question, answer in zip(questions, answers):
            cursor.execute(
                "INSERT INTO questions (subject_id, question_text, marks) VALUES (?, ?, ?)",
                (subject_id, question["text"], question["marks"])
            )
            question_id = cursor.lastrowid
            cursor.execute(
                "INSERT INTO answers (question_id, answer_text, word_count) VALUES (?, ?, ?)",
                (question_id, answer, len(answer.split()))
            )
        db.commit()
        db.close()

        return redirect(f"/subject/{subject_id}")

    return render_template("upload.html", subject_id=subject_id)

@app.route("/download/<int:subject_id>")
def download(subject_id):
    if "user_id" not in session:
        return redirect("/login")

    db = get_db()
    cursor = db.cursor()
    cursor.execute(
        """SELECT q.question_text, q.marks, a.answer_text
           FROM questions q
           LEFT JOIN answers a ON q.id = a.question_id
           WHERE q.subject_id=?""",
        (subject_id,)
    )
    rows = cursor.fetchall()
    cursor.execute("SELECT name FROM subjects WHERE id=?", (subject_id,))
    subject = cursor.fetchone()
    db.close()

    questions = [{"text": r[0], "marks": r[1]} for r in rows]
    answers = [r[2] or "" for r in rows]

    output_folder = f"outputs/{subject_id}"
    os.makedirs(output_folder, exist_ok=True)
    output_path = os.path.join(output_folder, "answers.pdf")

    generate_pdf(questions, answers, output_path, subject[0])

    return send_file(output_path, as_attachment=True,
                     download_name=f"{subject[0]}_answers.pdf")

if __name__ == "__main__":
    app.run(debug=True)
